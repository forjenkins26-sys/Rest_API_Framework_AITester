#!/usr/bin/env python3
"""Convert a test-plan markdown file to a .docx (Word) document.

Handles the markdown subset the test-plan template produces:
  - headings  # / ## / ###
  - paragraphs
  - bullet lists  - item   (one level of "  - " nesting)
  - blockquotes  > text
  - pipe tables  | a | b |  with a |---|---| separator row
  - horizontal rule  ---
  - inline  **bold**  and  _italic_

Usage:
    python md_to_docx.py <input.md> <output.docx>

Requires python-docx (`pip install python-docx`).
"""
import re
import sys

from docx import Document
from docx.shared import Pt


def add_runs(paragraph, text):
    """Add text to a paragraph, honoring **bold** and _italic_ inline markers."""
    # Split on bold/italic tokens while keeping the delimiters.
    for token in re.split(r"(\*\*.+?\*\*|_.+?_)", text):
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def is_table_sep(line):
    return bool(re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line)) and "-" in line


def split_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def main():
    if len(sys.argv) != 3:
        print("usage: md_to_docx.py <input.md> <output.docx>", file=sys.stderr)
        sys.exit(2)

    src, dst = sys.argv[1], sys.argv[2]
    with open(src, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Table block: a header row followed by a separator row.
        if stripped.startswith("|") and i + 1 < n and is_table_sep(lines[i + 1]):
            header = split_row(lines[i])
            i += 2  # skip header + separator
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Light Grid Accent 1"
            for c, text in enumerate(header):
                cell = table.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(text)
                run.bold = True
            for row in rows:
                cells = table.add_row().cells
                for c in range(len(header)):
                    cells[c].paragraphs[0].text = ""
                    add_runs(cells[c].paragraphs[0], row[c] if c < len(row) else "")
            doc.add_paragraph("")
            continue

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            doc.add_paragraph("").add_run().add_break()
            i += 1
            continue

        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, stripped[2:])
        elif re.match(r"^\s*[-*] ", line):
            indent = len(line) - len(line.lstrip())
            style = "List Bullet 2" if indent >= 2 else "List Bullet"
            p = doc.add_paragraph(style=style)
            add_runs(p, re.sub(r"^\s*[-*] ", "", line))
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s", "", stripped))
        else:
            p = doc.add_paragraph()
            add_runs(p, stripped)
        i += 1

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main()
